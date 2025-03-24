import os
import pdfplumber
import docx
import re
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
import requests
import json
import fitz
from django.shortcuts import redirect

def redirect_to_frontend(request):
    return redirect("/static/FE_2/index.html") 



GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
@csrf_exempt 
def upload_file(request):
    if request.method == "POST" and request.FILES.get("file"):
        uploaded_file = request.FILES["file"]

        # Define upload path
        upload_path = os.path.join(settings.MEDIA_ROOT, uploaded_file.name)

        # Save file
        with open(upload_path, "wb+") as destination:
            for chunk in uploaded_file.chunks():
                destination.write(chunk)

        return JsonResponse({"message": "File uploaded successfully!", "file_path": upload_path})

    return JsonResponse({"error": "No file uploaded"}, status=400)

@csrf_exempt
def chatbot_response(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            user_message = data.get("message", "")

            headers = {
                "Authorization": f"Bearer {settings.GROQ_API_KEY}",
                "Content-Type": "application/json",
            }

            payload = {
                "model": "llama3-8b-8192",  # Using LLaMA 3 model
                "messages": [{"role": "user", "content": user_message}],
                "temperature": 0.7
            }

            response = requests.post(GROQ_API_URL, headers=headers, json=payload)
            response_data = response.json()

            # Print full response to debug
            print("Groq API Full Response:", json.dumps(response_data, indent=4))

            # Extract bot reply (check if response format is valid)
            bot_reply = response_data.get("choices", [{}])[0].get("message", {}).get("content", None)

            if not bot_reply:
                return JsonResponse({"error": "Failed to generate a response", "api_response": response_data}, status=500)

            return JsonResponse({"bot_reply": bot_reply})

        except Exception as e:
            print("Error:", str(e))
            return JsonResponse({"error": str(e)}, status=500)

    return JsonResponse({"error": "Invalid request"}, status=400)

@csrf_exempt
def extract_toc(request):
    if request.method == "POST" and request.FILES.get("file"):
        uploaded_file = request.FILES["file"]
        file_extension = uploaded_file.name.split(".")[-1].lower()

        # Define the upload path
        upload_path = os.path.join(settings.MEDIA_ROOT, uploaded_file.name)
        with open(upload_path, "wb+") as destination:
            for chunk in uploaded_file.chunks():
                destination.write(chunk)

        # Extract TOC or Main Topics
        if file_extension == "pdf":
            toc = extract_toc_from_pdf(upload_path)
            if not toc:  # If no TOC found, extract main topics
                toc = extract_main_topics_from_pdf(upload_path)
        elif file_extension in ["doc", "docx"]:
            toc = extract_toc_from_docx(upload_path)
            if not toc:  # If no TOC found, extract main topics
                toc = extract_main_topics_from_docx(upload_path)
        else:
            return JsonResponse({"error": "Unsupported file format"}, status=400)

        return JsonResponse({"message": "TOC/Main Topics extracted successfully!", "toc": toc})

    return JsonResponse({"error": "No file uploaded"}, status=400)


def extract_toc_from_pdf(file_path):
    """ Extracts TOC from PDFs using regex if `get_toc()` fails. """
    doc = fitz.open(file_path)
    pdf_toc = doc.get_toc(simple=True)

    if pdf_toc:  # ✅ If `get_toc()` works, return structured TOC
        return [f"Page {entry[2]}: {entry[1]}" for entry in pdf_toc]

    # ❌ If no built-in TOC, extract using regex
    return extract_toc_by_regex(file_path)

def extract_toc_by_regex(file_path):
    """ Extracts TOC-like structures using regex from PDF text. """
    toc_list = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                toc_lines = re.findall(r'(.+?)\s+(\d+)', text)
                toc_list.extend([f"Page {num}: {title.strip()}" for title, num in toc_lines])
    return toc_list if toc_list else ["No TOC found, extracted main topics instead."]

def extract_toc_from_docx(file_path):
    doc = docx.Document(file_path)
    toc_list = []
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            toc_list.append(para.text)
    return toc_list




def extract_main_topics_from_pdf(file_path):
    """ Extracts main topics by detecting bold/large text from PDFs. """
    topics = set()
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            words = page.extract_words()
            if words:
                for word in words:
                    text_size = word.get("size", 0)  # Safely get size
                    font_type = word.get("fontname", "").lower()  # Detect bold text
                    if float(text_size) > 12 or "bold" in font_type:
                        topics.add(word["text"])
    return list(topics)


def extract_main_topics_from_docx(file_path):
    """ Extracts main topics from a DOCX by detecting headings or bold text. """
    doc = docx.Document(file_path)
    topics = []
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading") or any(run.bold for run in para.runs):
            topics.append(para.text)
    return topics

